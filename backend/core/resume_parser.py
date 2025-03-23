import os
import re
import tempfile
from typing import Dict, List, Optional, Tuple, Any

# For PDF parsing
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# For DOCX parsing
try:
    import docx
except ImportError:
    docx = None

# For OCR fallback
try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract = None
    Image = None


class ResumeParser:
    """Extract and structure content from resume files (PDF, DOCX)"""
    
    def __init__(self, ocr_enabled: bool = True):
        """
        Initialize resume parser
        
        Args:
            ocr_enabled: Whether to use OCR as fallback for unreadable PDFs
        """
        self.ocr_enabled = ocr_enabled and pytesseract is not None
        
        # Common section headers in resumes
        self.section_headers = {
            'education': ['education', 'academic background', 'qualifications', 'academic credentials'],
            'experience': ['experience', 'work experience', 'employment history', 'work history', 'professional experience'],
            'skills': ['skills', 'technical skills', 'core competencies', 'key skills', 'technical expertise'],
            'projects': ['projects', 'key projects', 'relevant projects'],
            'certifications': ['certifications', 'certificates', 'professional certifications'],
            'languages': ['languages', 'language proficiency'],
            'summary': ['summary', 'professional summary', 'profile', 'about me'],
            'contact': ['contact', 'contact information', 'personal information']
        }

    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse resume file and extract structured content
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary with extracted resume content
        """
        if not os.path.exists(file_path):
            return {"error": "File not found"}
            
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                if pdfplumber is None:
                    return {"error": "PDF parsing library not available. Install with: pip install pdfplumber"}
                text = self._extract_text_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                if docx is None:
                    return {"error": "DOCX parsing library not available. Install with: pip install python-docx"}
                text = self._extract_text_from_docx(file_path)
            else:
                return {"error": f"Unsupported file format: {file_ext}"}
                
            if not text.strip():
                return {"error": "No text content extracted from file"}
                
            # Process and structure the extracted text
            structured_data = self._structure_resume_content(text)
            return structured_data
            
        except Exception as e:
            return {"error": f"Error parsing resume: {str(e)}"}

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text content from PDF files"""
        full_text = ""
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    full_text += page_text + "\n"
                    
            # If text extraction failed or returned minimal content, try OCR
            if len(full_text.strip()) < 100 and self.ocr_enabled:
                full_text = self._ocr_fallback(file_path)
                
            return full_text
        except Exception as e:
            print(f"PDF extraction error: {str(e)}")
            if self.ocr_enabled:
                return self._ocr_fallback(file_path)
            return ""

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text content from DOCX files"""
        full_text = ""
        
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                full_text += para.text + "\n"
                
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + " "
                    full_text += "\n"
                    
            return full_text
        except Exception as e:
            print(f"DOCX extraction error: {str(e)}")
            return ""

    def _ocr_fallback(self, file_path: str) -> str:
        """Use OCR as fallback for documents that can't be parsed directly"""
        if not self.ocr_enabled:
            return ""
            
        full_text = ""
        
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                # Convert PDF pages to images
                from pdf2image import convert_from_path
                images = convert_from_path(file_path)
                
                # Process each page with OCR
                for i, image in enumerate(images):
                    image_path = os.path.join(temp_dir, f'page_{i}.png')
                    image.save(image_path, 'PNG')
                    
                    # Extract text with OCR
                    page_text = pytesseract.image_to_string(Image.open(image_path))
                    full_text += page_text + "\n"
                    
            return full_text
        except Exception as e:
            print(f"OCR fallback error: {str(e)}")
            return ""

    def _structure_resume_content(self, text: str) -> Dict[str, Any]:
        """
        Process raw text and structure into resume sections
        
        Args:
            text: Raw text extracted from resume
            
        Returns:
            Dictionary with structured resume content
        """
        # Initialize structured data
        structured_data = {
            "raw_text": text,
            "contact_info": self._extract_contact_info(text),
            "sections": {},
            "metadata": {
                "has_education": False,
                "has_experience": False,
                "has_skills": False,
                "extracted_skills": [],
                "word_count": len(text.split()),
                "line_count": len(text.splitlines())
            }
        }
        
        # Split text into lines and process by sections
        lines = text.splitlines()
        current_section = "unknown"
        section_content = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if this line is a section header
            detected_section = self._detect_section(line)
            if detected_section:
                # Save content for previous section
                if current_section != "unknown" and current_section not in section_content:
                    section_content[current_section] = []
                
                current_section = detected_section
                structured_data["metadata"][f"has_{current_section}"] = True
            else:
                # Add line to current section
                if current_section not in section_content:
                    section_content[current_section] = []
                section_content[current_section].append(line)
        
        # Process each section with specific extractors
        structured_data["sections"] = section_content
        
        # Extract skills
        structured_data["metadata"]["extracted_skills"] = self._extract_skills(text)
        
        # Extract education
        structured_data["education"] = self._extract_education(section_content.get("education", []))
        
        # Extract experience
        structured_data["experience"] = self._extract_experience(section_content.get("experience", []))
        
        return structured_data

    def _detect_section(self, line: str) -> Optional[str]:
        """Detect if a line is a section header and return section type"""
        line_lower = line.lower()
        
        for section, headers in self.section_headers.items():
            for header in headers:
                # Check exact match or header ending with colon
                if line_lower == header or line_lower == f"{header}:" or \
                   line_lower.startswith(f"{header}:") or line_lower.startswith(f"{header} "):
                    return section
                    
        return None

    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information from resume text"""
        contact_info = {
            "email": "",
            "phone": "",
            "linkedin": "",
            "github": "",
            "website": ""
        }
        
        # Extract email
        email_pattern = r'[\w.+-]+@[\w-]+\.[\w.-]+'
        email_matches = re.findall(email_pattern, text)
        if email_matches:
            contact_info["email"] = email_matches[0]
        
        # Extract phone number (various formats)
        phone_pattern = r'(?:\+\d{1,3}[- ]?)?\(?\d{3}\)?[- ]?\d{3}[- ]?\d{4}'
        phone_matches = re.findall(phone_pattern, text)
        if phone_matches:
            contact_info["phone"] = phone_matches[0]
        
        # Extract LinkedIn URL
        linkedin_pattern = r'linkedin\.com/(?:in|profile)/[\w-]+'
        linkedin_matches = re.findall(linkedin_pattern, text.lower())
        if linkedin_matches:
            contact_info["linkedin"] = "https://" + linkedin_matches[0]
        
        # Extract GitHub URL
        github_pattern = r'github\.com/[\w-]+'
        github_matches = re.findall(github_pattern, text.lower())
        if github_matches:
            contact_info["github"] = "https://" + github_matches[0]
        
        # Extract website/portfolio
        website_pattern = r'(?:https?://)?(?:www\.)?[\w-]+\.[\w.-]+(?:/[\w.-]*)*/?'
        website_matches = re.findall(website_pattern, text.lower())
        
        for url in website_matches:
            if "linkedin" not in url and "github" not in url:
                if not url.startswith(('http://', 'https://')):
                    url = 'https://' + url
                contact_info["website"] = url
                break
        
        return contact_info

    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text"""
        # Common technical skills to look for
        common_skills = [
            # Programming Languages
            "python", "java", "javascript", "c++", "c#", "ruby", "php", "swift", "kotlin", "golang",
            # Web Technologies
            "html", "css", "react", "angular", "vue", "node.js", "django", "flask", "spring", "express",
            # Data Science
            "machine learning", "data science", "tensorflow", "pytorch", "pandas", "numpy", "scikit-learn",
            # Database
            "sql", "mysql", "postgresql", "mongodb", "oracle", "firebase", "dynamodb", "redis",
            # DevOps
            "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "git", "ci/cd", "terraform",
            # General Skills
            "leadership", "project management", "agile", "scrum", "communication"
        ]
        
        # Extract skills from text
        found_skills = []
        text_lower = text.lower()
        
        for skill in common_skills:
            # Look for skill as whole word with word boundaries
            if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
                found_skills.append(skill)
        
        return found_skills

    def _extract_education(self, education_lines: List[str]) -> List[Dict[str, str]]:
        """Extract education information from education section"""
        education_entries = []
        
        if not education_lines:
            return education_entries
        
        # Join lines to process chunks of text
        education_text = " ".join(education_lines)
        
        # Common degree patterns
        degree_patterns = [
            r"(Bachelor(?:'s)?|Master(?:'s)?|Ph\.D\.|Doctorate|B\.S\.|M\.S\.|B\.A\.|M\.A\.|B\.E\.|M\.E\.|M\.B\.A\.|B\.Tech|M\.Tech)",
            r"(Bachelor of [A-Za-z\s]+|Master of [A-Za-z\s]+|Doctor of [A-Za-z\s]+)"
        ]
        
        # Look for degree mentions and split education section
        degrees = []
        for pattern in degree_patterns:
            degrees.extend(re.findall(pattern, education_text, re.IGNORECASE))
        
        # If we found degrees, try to extract information around them
        for degree in degrees:
            # Find the text chunk around this degree (50 chars before and after)
            degree_idx = education_text.lower().find(degree.lower())
            if degree_idx != -1:
                start_idx = max(0, degree_idx - 50)
                end_idx = min(len(education_text), degree_idx + len(degree) + 50)
                chunk = education_text[start_idx:end_idx]
                
                # Extract university/institution
                university = ""
                university_pattern = r"(?:at|from|University of|College|Institute|School of) ([A-Za-z\s]+)"
                university_match = re.search(university_pattern, chunk, re.IGNORECASE)
                if university_match:
                    university = university_match.group(1).strip()
                
                # Extract dates
                date_pattern = r"(\d{4}\s*(?:-|–|to)\s*(?:\d{4}|present|current)|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{4})"
                date_match = re.search(date_pattern, chunk, re.IGNORECASE)
                date = date_match.group(1) if date_match else ""
                
                # Extract GPA if available
                gpa = ""
                gpa_pattern = r"GPA:?\s*(\d+\.\d+|[A-Z]+[-+]?)"
                gpa_match = re.search(gpa_pattern, chunk, re.IGNORECASE)
                if gpa_match:
                    gpa = gpa_match.group(1)
                
                education_entries.append({
                    "degree": degree.strip(),
                    "institution": university,
                    "date": date.strip(),
                    "gpa": gpa
                })
        
        return education_entries

    def _extract_experience(self, experience_lines: List[str]) -> List[Dict[str, str]]:
        """Extract work experience information from experience section"""
        experience_entries = []
        
        if not experience_lines:
            return experience_entries
        
        # Join lines for processing
        experience_text = "\n".join(experience_lines)
        
        # Split into potential job entries
        # Look for patterns like "Company Name | Job Title | Date Range"
        job_blocks = []
        
        # Pattern matching for job entries (company followed by title, or title followed by company)
        company_patterns = [
            r"([A-Z][A-Za-z0-9\s&.,]+)(?:\s*[-|•]\s*|\s+at\s+|\s+in\s+)([A-Za-z\s]+)(?:\s*[-|•]\s*|\s+)(\d{4}\s*(?:-|–|to)\s*(?:\d{4}|present|current))",
            r"([A-Za-z\s]+)(?:\s+at\s+|\s+in\s+)([A-Z][A-Za-z0-9\s&.,]+)(?:\s*[-|•]\s*|\s+)(\d{4}\s*(?:-|–|to)\s*(?:\d{4}|present|current))"
        ]
        
        for pattern in company_patterns:
            matches = re.findall(pattern, experience_text)
            for match in matches:
                if len(match) >= 3:
                    # Try to determine if the first or second group is the company
                    if match[0][0].isupper() and len(match[0].split()) <= 5:
                        company, title, date = match
                    else:
                        title, company, date = match
                    
                    experience_entries.append({
                        "title": title.strip(),
                        "company": company.strip(),
                        "date": date.strip(),
                        "description": ""  # Will populate this in the next step
                    })
        
        # If we couldn't find structured entries, try simpler approach
        if not experience_entries:
            # Split by multiple newlines which often separate job entries
            potential_entries = re.split(r'\n\s*\n', experience_text)
            
            for entry in potential_entries:
                if len(entry.split()) > 5:  # Skip entries too short to be jobs
                    lines = entry.split('\n')
                    if not lines:
                        continue
                    
                    # Try to extract potential job title and company from first few lines
                    potential_title = ""
                    potential_company = ""
                    potential_date = ""
                    
                    for i, line in enumerate(lines[:3]):
                        # Look for dates in this line
                        date_match = re.search(r'(\d{4}\s*(?:-|–|to)\s*(?:\d{4}|present|current))', line)
                        if date_match:
                            potential_date = date_match.group(1)
                            
                            # If we found a date, the title/company are likely in this or previous line
                            title_company = line.replace(potential_date, "").strip()
                            
                            # Split by common separators
                            parts = re.split(r'[|•-]', title_company)
                            if len(parts) >= 2:
                                potential_title = parts[0].strip()
                                potential_company = parts[1].strip()
                            break
                            
                    # Build description from remaining lines
                    description = "\n".join(lines[i+1:] if i < len(lines) else [])
                    
                    if potential_title or potential_company:
                        experience_entries.append({
                            "title": potential_title,
                            "company": potential_company,
                            "date": potential_date,
                            "description": description.strip()
                        })
        
        return experience_entries 