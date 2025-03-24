"""
Resume Parser Module

This module extracts structured information from resume files in various formats,
identifying key elements such as contact information, skills, work experience, and education.
"""

import os
import re
import json
import tempfile
from typing import Dict, List, Any, Optional, Tuple, Union
from datetime import datetime
import logging

# Import settings
from config.settings import BASE_DIR

# Create logger
logger = logging.getLogger(__name__)

# Try importing document parsing libraries
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Install with: pip install python-docx")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logger.warning("PyPDF2 not available. Install with: pip install PyPDF2")

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False
    logger.warning("pdfminer.six not available. Install with: pip install pdfminer.six")

# Try importing NLP libraries
try:
    import spacy
    import nltk
    from nltk.tokenize import word_tokenize
    from nltk.corpus import stopwords
    
    # Initialize NLTK resources
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt', quiet=True)
    
    try:
        nltk.data.find('corpora/stopwords')
    except LookupError:
        nltk.download('stopwords', quiet=True)
    
    try:
        nltk.data.find('taggers/averaged_perceptron_tagger')
    except LookupError:
        nltk.download('averaged_perceptron_tagger', quiet=True)
    
    try:
        nltk.data.find('chunkers/maxent_ne_chunker')
    except LookupError:
        nltk.download('maxent_ne_chunker', quiet=True)
    
    try:
        nltk.data.find('corpora/words')
    except LookupError:
        nltk.download('words', quiet=True)
    
    NLP_AVAILABLE = True
    
    # Load spaCy model
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        logger.warning("spaCy model not found. Run: python -m spacy download en_core_web_sm")
        NLP_AVAILABLE = False
except ImportError:
    NLP_AVAILABLE = False
    logger.warning("Advanced NLP libraries not available. Install with: pip install spacy nltk")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a .docx file
    
    Args:
        file_path: Path to .docx file
        
    Returns:
        str: Extracted text
    """
    if not DOCX_AVAILABLE:
        logger.error("python-docx not available. Cannot extract from .docx")
        return ""
    
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from docx: {e}")
        return ""


def extract_text_from_pdf_pypdf2(file_path: str) -> str:
    """
    Extract text from a PDF file using PyPDF2
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        str: Extracted text
    """
    if not PYPDF2_AVAILABLE:
        logger.error("PyPDF2 not available. Cannot extract from PDF")
        return ""
    
    try:
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF using PyPDF2: {e}")
        return ""


def extract_text_from_pdf_pdfminer(file_path: str) -> str:
    """
    Extract text from a PDF file using pdfminer
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        str: Extracted text
    """
    if not PDFMINER_AVAILABLE:
        logger.error("pdfminer.six not available. Cannot extract from PDF")
        return ""
    
    try:
        return pdf_extract_text(file_path)
    except Exception as e:
        logger.error(f"Error extracting text from PDF using pdfminer: {e}")
        return ""


def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from a .txt file
    
    Args:
        file_path: Path to text file
        
    Returns:
        str: Extracted text
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    except Exception as e:
        logger.error(f"Error extracting text from txt: {e}")
        return ""


def extract_text_from_resume(file_path: str) -> str:
    """
    Extract text from a resume file based on its extension
    
    Args:
        file_path: Path to resume file
        
    Returns:
        str: Extracted text
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.pdf':
        # Try pdfminer first, then fall back to PyPDF2
        if PDFMINER_AVAILABLE:
            text = extract_text_from_pdf_pdfminer(file_path)
            if text.strip():
                return text
        
        if PYPDF2_AVAILABLE:
            return extract_text_from_pdf_pypdf2(file_path)
            
        return ""
        
    elif file_ext == '.docx':
        return extract_text_from_docx(file_path)
        
    elif file_ext == '.txt':
        return extract_text_from_txt(file_path)
        
    else:
        logger.error(f"Unsupported file format: {file_ext}")
        return ""


def preprocess_text(text: str) -> str:
    """
    Preprocess resume text to improve parsing accuracy
    
    Args:
        text: Raw resume text
        
    Returns:
        str: Preprocessed text
    """
    # Replace multiple newlines with a single newline
    text = re.sub(r'\n+', '\n', text)
    
    # Replace multiple spaces with a single space
    text = re.sub(r'\s+', ' ', text)
    
    # Replace pipe characters (common in table extraction)
    text = re.sub(r'\s*\|\s*', ' | ', text)
    
    return text.strip()


def extract_sections(text: str) -> Dict[str, str]:
    """
    Extract main sections from resume text
    
    Args:
        text: Resume text
        
    Returns:
        dict: Extracted sections
    """
    # Define section patterns
    section_patterns = {
        "contact": r"(?i)(CONTACT|PERSONAL|DETAILS|INFORMATION|PROFILE)",
        "summary": r"(?i)(SUMMARY|OBJECTIVE|PROFESSIONAL\s+SUMMARY|CAREER\s+OBJECTIVE|ABOUT|PROFILE)",
        "education": r"(?i)(EDUCATION|ACADEMIC|QUALIFICATION|DEGREE|UNIVERSITY)",
        "experience": r"(?i)(EXPERIENCE|EMPLOYMENT|WORK|CAREER|JOB|PROFESSIONAL\s+EXPERIENCE)",
        "skills": r"(?i)(SKILLS|ABILITIES|COMPETENCIES|TECHNICAL|EXPERTISE|PROFICIENCIES)",
        "projects": r"(?i)(PROJECTS|PROJECT\s+EXPERIENCE|ASSIGNMENTS)",
        "certifications": r"(?i)(CERTIFICATIONS|CERTIFICATES|ACCREDITATIONS|COURSES)",
        "languages": r"(?i)(LANGUAGES|LANGUAGE\s+PROFICIENCY)",
        "awards": r"(?i)(AWARDS|ACHIEVEMENTS|HONORS|RECOGNITION)",
        "interests": r"(?i)(INTERESTS|HOBBIES|ACTIVITIES|VOLUNTEER)"
    }
    
    # Split text into lines
    lines = text.split('\n')
    
    # Identify section boundaries
    sections = {}
    current_section = None
    section_content = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        found_section = False
        for section_name, pattern in section_patterns.items():
            if re.match(pattern, line):
                # If we were building a previous section, save it
                if current_section:
                    sections[current_section] = '\n'.join(section_content)
                
                # Start a new section
                current_section = section_name
                section_content = []
                found_section = True
                break
        
        if not found_section and current_section:
            section_content.append(line)
    
    # Add the last section
    if current_section and section_content:
        sections[current_section] = '\n'.join(section_content)
    
    return sections


def extract_contact_info(text: str) -> Dict[str, Any]:
    """
    Extract contact information from resume text
    
    Args:
        text: Resume text
        
    Returns:
        dict: Extracted contact information
    """
    contact_info = {
        "name": None,
        "email": None,
        "phone": None,
        "address": None,
        "linkedin": None,
        "website": None
    }
    
    # Email pattern
    email_pattern = r'[\w.+-]+@[\w-]+\.[\w.-]+'
    email_matches = re.findall(email_pattern, text)
    if email_matches:
        contact_info["email"] = email_matches[0]
    
    # Phone pattern - various formats
    phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    phone_matches = re.findall(phone_pattern, text)
    if phone_matches:
        contact_info["phone"] = phone_matches[0]
    
    # LinkedIn URL
    linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/profile/view\?id=)[\w-]+'
    linkedin_matches = re.findall(linkedin_pattern, text.lower())
    if linkedin_matches:
        contact_info["linkedin"] = "https://www." + linkedin_matches[0]
    
    # Website URL
    website_pattern = r'(?:https?://)?(?:www\.)?([a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)'
    website_matches = re.findall(website_pattern, text)
    if website_matches:
        # Filter out common domains like linkedin, facebook, etc.
        excluded_domains = ['linkedin.com', 'facebook.com', 'twitter.com', 'instagram.com']
        for match in website_matches:
            if not any(domain in match.lower() for domain in excluded_domains):
                contact_info["website"] = match
                break
    
    # Name (best guess from first few lines)
    lines = text.split('\n')
    for line in lines[:5]:  # Check first 5 lines
        # Name is often the first line with 2-4 words and no special characters
        words = line.strip().split()
        if 2 <= len(words) <= 4 and re.match(r'^[A-Za-z\s\.-]+$', line) and len(line) < 40:
            contact_info["name"] = line.strip()
            break
    
    return contact_info


def extract_education(text: str) -> List[Dict[str, str]]:
    """
    Extract education information from resume text
    
    Args:
        text: Resume text with education section
        
    Returns:
        list: List of education entries
    """
    education_list = []
    
    # Split into education entries (separated by years or degrees)
    degree_patterns = [
        r'(?i)(B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?A\.?|Ph\.?D\.?|MBA|Bachelor|Master|Doctor|Associate)',
        r'(?i)(Engineering|Business|Science|Arts|Administration|Technology)'
    ]
    year_pattern = r'(19|20)\d{2}(\s*(-|to|–)\s*(19|20)\d{2}|present|current|\s*-\s*\d{2})?'
    
    # Find year matches to split the text
    year_matches = list(re.finditer(year_pattern, text))
    
    if year_matches:
        # Use years to split into education entries
        entries = []
        for i in range(len(year_matches)):
            start = year_matches[i].start()
            
            # Find the start of the previous entry or the beginning of text
            prev_end = 0 if i == 0 else year_matches[i-1].end()
            
            # Find where this entry likely starts (after previous match or at beginning)
            # Look backward for the beginning of the line
            entry_start = text.rfind('\n', prev_end, start)
            if entry_start == -1:
                entry_start = prev_end
            
            # Calculate end point (next match start or end of text)
            end = len(text) if i == len(year_matches) - 1 else year_matches[i+1].start()
            
            # Find the actual end - look for the next newline after the year
            entry_end = text.find('\n', year_matches[i].end(), end)
            if entry_end == -1:
                entry_end = end
            
            # Extract the entry
            entry = text[entry_start:entry_end].strip()
            if entry:
                entries.append(entry)
    else:
        # If no years found, try splitting by degree patterns
        entries = []
        for pattern in degree_patterns:
            degree_matches = list(re.finditer(pattern, text))
            
            if degree_matches:
                for i in range(len(degree_matches)):
                    start = degree_matches[i].start()
                    
                    # Find the start of the previous entry or the beginning of text
                    prev_end = 0 if i == 0 else degree_matches[i-1].end()
                    
                    # Find where this entry likely starts
                    entry_start = text.rfind('\n', prev_end, start)
                    if entry_start == -1:
                        entry_start = prev_end
                    
                    # Calculate end point
                    end = len(text) if i == len(degree_matches) - 1 else degree_matches[i+1].start()
                    
                    # Find the actual end
                    entry_end = text.find('\n', degree_matches[i].end(), end)
                    if entry_end == -1:
                        entry_end = end
                    
                    # Extract the entry
                    entry = text[entry_start:entry_end].strip()
                    if entry:
                        entries.append(entry)
                
                if entries:
                    break  # If we found entries with one pattern, stop
    
    # If still no entries, just split by newlines
    if not entries:
        entries = [line.strip() for line in text.split('\n') if line.strip()]
    
    # Process each entry
    for entry in entries:
        education_item = {"institution": None, "degree": None, "year": None, "description": None}
        
        # Extract year
        year_match = re.search(year_pattern, entry)
        if year_match:
            education_item["year"] = year_match.group(0).strip()
            
            # Remove year from entry to simplify further extraction
            entry = entry.replace(year_match.group(0), "").strip()
        
        # Extract degree
        degree_match = None
        for pattern in degree_patterns:
            match = re.search(pattern, entry, re.IGNORECASE)
            if match:
                degree_match = match
                break
                
        if degree_match:
            # Look for the complete degree (e.g., "Bachelor of Science in Computer Science")
            # by extracting text around the match
            start = max(0, degree_match.start() - 20)
            end = min(len(entry), degree_match.end() + 50)
            degree_area = entry[start:end]
            
            # Look for degree phrases
            degree_phrase_pattern = r'(?i)(Bachelor|Master|Doctor|Ph\.?D\.?|MBA|B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?)(\s+of\s+)?([\w\s]+)(\s+in\s+)([\w\s]+)'
            degree_phrase_match = re.search(degree_phrase_pattern, degree_area)
            
            if degree_phrase_match:
                education_item["degree"] = degree_phrase_match.group(0).strip()
            else:
                # Just use the basic match
                education_item["degree"] = degree_match.group(0).strip()
            
            # Remove degree from entry
            entry = entry.replace(education_item["degree"], "").strip()
        
        # The remaining text is likely to contain the institution and possibly description
        # For institution, look for common words like "University", "College", "Institute"
        institution_pattern = r'(?i)(University|College|Institute|School) of ([\w\s]+)|([\w\s]+) (University|College|Institute|School)'
        institution_match = re.search(institution_pattern, entry)
        
        if institution_match:
            education_item["institution"] = institution_match.group(0).strip()
            
            # Remove institution from entry
            entry = entry.replace(education_item["institution"], "").strip()
        else:
            # If no clear institution pattern, the first part is likely the institution
            lines = entry.split('\n')
            if lines:
                education_item["institution"] = lines[0].strip()
                entry = entry.replace(lines[0], "").strip()
        
        # Remaining text is the description
        if entry:
            education_item["description"] = entry.strip()
        
        # Only add if we have at least institution or degree
        if education_item["institution"] or education_item["degree"]:
            education_list.append(education_item)
    
    return education_list


def extract_experience(text: str) -> List[Dict[str, str]]:
    """
    Extract work experience information from resume text
    
    Args:
        text: Resume text with experience section
        
    Returns:
        list: List of experience entries
    """
    experience_list = []
    
    # Look for date patterns to identify experience entries
    date_pattern = r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\s*(-|to|–)\s*(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}|(\d{4})\s*(-|to|–)\s*(\d{4}|Present|Current|Now)'
    
    # Find all date ranges
    date_matches = list(re.finditer(date_pattern, text, re.IGNORECASE))
    
    if date_matches:
        # Use date ranges to split into experience entries
        entries = []
        for i in range(len(date_matches)):
            start = date_matches[i].start()
            
            # Find the start of the previous entry or the beginning of text
            prev_end = 0 if i == 0 else date_matches[i-1].end()
            
            # Find where this entry likely starts (after previous match or at beginning)
            # Look backward for the beginning of the line
            entry_start = text.rfind('\n', prev_end, start)
            if entry_start == -1:
                entry_start = prev_end
            
            # Calculate end point (next match start or end of text)
            end = len(text) if i == len(date_matches) - 1 else date_matches[i+1].start()
            
            # Get the text slice that likely represents this entry
            entry_text = text[entry_start:end].strip()
            
            if entry_text:
                # Split into lines to find job title and company
                lines = entry_text.split('\n')
                
                # Extract date range
                date_range = date_matches[i].group(0)
                
                # First lines before the date are likely job title and company
                job_title = None
                company = None
                
                # Look for the line containing the date range
                date_line_idx = -1
                for idx, line in enumerate(lines):
                    if date_range in line:
                        date_line_idx = idx
                        break
                
                # Extract job title and company based on position of date
                if date_line_idx >= 0:
                    # Date is on the same line as either job title or company
                    if date_line_idx == 0:
                        # Date is on the first line, it may contain job title or company
                        job_title_line = lines[0].replace(date_range, "").strip()
                        
                        # Check if there's a second line that might be the company
                        if len(lines) > 1:
                            company = lines[1].strip()
                            job_title = job_title_line
                        else:
                            # Only one line, use heuristics to decide if it's job title or company
                            # Companies often have Inc, LLC, Ltd, etc.
                            if re.search(r'(?i)(Inc|LLC|Ltd|Limited|Corporation|Corp|Company)', job_title_line):
                                company = job_title_line
                            else:
                                job_title = job_title_line
                    else:
                        # Date is not on first line, so first line is likely job title
                        job_title = lines[0].strip()
                        # Line with date is likely company (remove date part)
                        company = lines[date_line_idx].replace(date_range, "").strip()
                else:
                    # Date was not found in lines (should not happen due to our regex search)
                    # Fallback: assume first line is job title, second is company
                    if len(lines) > 0:
                        job_title = lines[0].strip()
                    if len(lines) > 1:
                        company = lines[1].strip()
                
                # Description is everything after company line
                description_start = date_line_idx + 1 if date_line_idx >= 0 else 2
                description = '\n'.join(lines[description_start:]).strip()
                
                experience_list.append({
                    "job_title": job_title,
                    "company": company,
                    "date_range": date_range,
                    "description": description
                })
    else:
        # If no date patterns, try another approach: look for job title patterns
        job_title_pattern = r'(?i)(Engineer|Manager|Developer|Director|Consultant|Analyst|Designer|Specialist|Coordinator|Assistant|Lead|Head)'
        
        job_matches = list(re.finditer(job_title_pattern, text))
        
        if job_matches:
            # Segment by job titles
            entries = []
            for i in range(len(job_matches)):
                start = job_matches[i].start()
                
                # Find the start of the entry
                entry_start = text.rfind('\n', 0, start)
                if entry_start == -1:
                    entry_start = 0
                
                # Calculate end point
                end = len(text) if i == len(job_matches) - 1 else job_matches[i+1].start()
                
                # Extract the entry
                entry = text[entry_start:end].strip()
                if entry:
                    # Basic parsing: assume first line is job title, second is company
                    lines = entry.split('\n')
                    job_title = lines[0].strip() if lines else None
                    company = lines[1].strip() if len(lines) > 1 else None
                    description = '\n'.join(lines[2:]).strip() if len(lines) > 2 else None
                    
                    # Try to extract date range with a more relaxed pattern
                    date_range = None
                    date_pattern_simple = r'\d{4}\s*(-|to|–)\s*(\d{4}|Present|Current|Now)'
                    date_match = re.search(date_pattern_simple, entry)
                    if date_match:
                        date_range = date_match.group(0)
                    
                    experience_list.append({
                        "job_title": job_title,
                        "company": company,
                        "date_range": date_range,
                        "description": description
                    })
    
    # Clean up the experience items
    for item in experience_list:
        # Make sure we don't have any None values
        for key in item:
            if item[key] is None:
                item[key] = ""
                
        # Remove date range from title or company if it's still there
        if item["date_range"]:
            item["job_title"] = item["job_title"].replace(item["date_range"], "").strip()
            item["company"] = item["company"].replace(item["date_range"], "").strip()
    
    return experience_list


def extract_skills(text: str) -> List[str]:
    """
    Extract skills from resume text
    
    Args:
        text: Resume text
        
    Returns:
        list: List of skills
    """
    # Common skill keyword lists
    technical_skills = [
        # Programming languages
        "Python", "Java", "JavaScript", "C\\+\\+", "C#", "PHP", "Ruby", "Swift", "Kotlin", "Go", 
        "Rust", "TypeScript", "Scala", "R", "MATLAB", "Perl", "Shell", "Bash", "SQL",
        
        # Frameworks and libraries
        "React", "Angular", "Vue", "Django", "Flask", "Spring", "ASP\\.NET", "Laravel", "Ruby on Rails",
        "TensorFlow", "PyTorch", "Keras", "Scikit-learn", "Pandas", "NumPy", "Express", "Node\\.js",
        
        # Database
        "MySQL", "PostgreSQL", "MongoDB", "Oracle", "SQL Server", "SQLite", "Redis", "Cassandra",
        "DynamoDB", "Firebase", "Elasticsearch",
        
        # Cloud & DevOps
        "AWS", "Azure", "Google Cloud", "Docker", "Kubernetes", "Jenkins", "GitLab CI", "Travis CI",
        "Terraform", "Ansible", "Puppet", "Chef", "CI/CD", "DevOps",
        
        # Web technologies
        "HTML", "CSS", "SASS", "LESS", "Bootstrap", "Material UI", "REST API", "GraphQL", "JSON",
        "XML", "SOAP", "WebSockets",
        
        # Mobile
        "Android", "iOS", "React Native", "Flutter", "Xamarin", "Swift", "Objective-C",
        
        # Tools & Software
        "Git", "SVN", "Jira", "Confluence", "Trello", "Slack", "Microsoft Office", "Excel", "PowerPoint",
        "Photoshop", "Illustrator", "Figma", "Sketch", "Adobe XD", "Tableau", "Power BI"
    ]
    
    soft_skills = [
        "Communication", "Leadership", "Teamwork", "Problem solving", "Critical thinking", 
        "Decision making", "Time management", "Adaptability", "Flexibility", "Creativity",
        "Emotional intelligence", "Conflict resolution", "Negotiation", "Presentation",
        "Public speaking", "Writing", "Customer service", "Project management"
    ]
    
    all_skills = technical_skills + soft_skills
    
    # Create a pattern to find skills
    # We'll look for them as standalone words or phrases to avoid false positives
    skill_pattern = r'\b(' + '|'.join(all_skills) + r')\b'
    
    # Find all skill matches
    skill_matches = re.findall(skill_pattern, text, re.IGNORECASE)
    
    # Remove duplicates while preserving order
    seen = set()
    skills = []
    for skill in skill_matches:
        skill_lower = skill.lower()
        if skill_lower not in seen:
            # Use the original case from the all_skills list if possible
            original_case = next((s for s in all_skills if s.lower() == skill_lower), skill)
            skills.append(original_case)
            seen.add(skill_lower)
    
    return skills


def extract_entities_with_nlp(text: str) -> Dict[str, List[str]]:
    """
    Extract entities using NLP libraries
    
    Args:
        text: Resume text
        
    Returns:
        dict: Extracted entities by type
    """
    entities = {
        "skills": [],
        "job_titles": [],
        "companies": [],
        "educational_institutions": [],
        "degrees": [],
        "certifications": []
    }
    
    if not NLP_AVAILABLE:
        return entities
    
    try:
        # Process with spaCy
        doc = nlp(text)
        
        # Extract named entities
        for ent in doc.ents:
            if ent.label_ == "ORG":
                # Could be a company or educational institution
                org_name = ent.text.strip()
                
                # Check if it's likely an educational institution
                if re.search(r'(?i)university|college|institute|school', org_name):
                    entities["educational_institutions"].append(org_name)
                else:
                    entities["companies"].append(org_name)
            
            elif ent.label_ == "PERSON":
                # Could be a person's name, not much use for resume parsing
                pass
                
            elif ent.label_ == "DATE":
                # We've already handled dates in other functions
                pass
        
        # Extract potential job titles using POS patterns
        job_title_patterns = [
            [{"POS": "ADJ"}, {"POS": "NOUN"}],  # Senior Engineer
            [{"POS": "NOUN"}, {"POS": "NOUN"}],  # Software Engineer
            [{"POS": "PROPN"}, {"POS": "NOUN"}]  # Python Developer
        ]
        
        for pattern in job_title_patterns:
            matcher = spacy.matcher.Matcher(nlp.vocab)
            matcher.add("JOB_TITLE", [pattern])
            matches = matcher(doc)
            
            for match_id, start, end in matches:
                span = doc[start:end]
                # Filter to likely job titles
                if any(term in span.text.lower() for term in ["engineer", "developer", "manager", "analyst", "designer"]):
                    entities["job_titles"].append(span.text)
        
        # Extract degrees and certifications using patterns
        degree_patterns = [
            r'(?i)(Bachelor|Master|Doctor|Ph\.?D\.?|MBA)(\s+of\s+)?([\w\s]+)(\s+in\s+)([\w\s]+)',
            r'(?i)(B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?) in ([\w\s]+)'
        ]
        
        cert_patterns = [
            r'(?i)(Certified|Certificate in|Certification in) ([\w\s]+)',
            r'(?i)([\w\s]+) (Certification|Certificate)'
        ]
        
        for pattern in degree_patterns:
            degrees = re.findall(pattern, text)
            for degree in degrees:
                if isinstance(degree, tuple):
                    entities["degrees"].append(" ".join(degree).strip())
                else:
                    entities["degrees"].append(degree.strip())
        
        for pattern in cert_patterns:
            certs = re.findall(pattern, text)
            for cert in certs:
                if isinstance(cert, tuple):
                    entities["certifications"].append(" ".join(cert).strip())
                else:
                    entities["certifications"].append(cert.strip())
        
        # Add skills that might not be in our predefined list
        # Look for noun chunks that might be skills
        for chunk in doc.noun_chunks:
            # Skills are often 1-3 word technical terms
            if 1 <= len(chunk) <= 3 and not chunk.text.lower() in ["i", "me", "my", "he", "she", "they", "we"]:
                # Check if it's near skill-indicating words
                skill_context = doc[max(0, chunk.start - 5):min(len(doc), chunk.end + 5)].text.lower()
                if any(word in skill_context for word in ["skill", "proficient", "experience", "knowledge", "familiar"]):
                    entities["skills"].append(chunk.text)
    
    except Exception as e:
        logger.error(f"Error in NLP entity extraction: {e}")
    
    # Remove duplicates
    for entity_type in entities:
        entities[entity_type] = list(set(entities[entity_type]))
    
    return entities


def parse_resume(file_path: str) -> Dict[str, Any]:
    """
    Parse resume file into structured data
    
    Args:
        file_path: Path to resume file
        
    Returns:
        dict: Structured resume data
    """
    # Initialize return structure
    parsed_data = {
        "file_path": file_path,
        "file_type": os.path.splitext(file_path)[1].lower(),
        "raw_text": "",
        "sections": {},
        "contact_info": {},
        "education": [],
        "experience": [],
        "skills": [],
        "entities": {},
        "metadata": {
            "parsed_at": datetime.now().isoformat(),
            "parser_version": "1.0",
            "word_count": 0
        }
    }
    
    try:
        # Extract text from file
        raw_text = extract_text_from_resume(file_path)
        if not raw_text:
            parsed_data["error"] = "Failed to extract text from file"
            return parsed_data
        
        parsed_data["raw_text"] = raw_text
        parsed_data["metadata"]["word_count"] = len(raw_text.split())
        
        # Preprocess text
        processed_text = preprocess_text(raw_text)
        
        # Extract sections
        sections = extract_sections(processed_text)
        parsed_data["sections"] = sections
        
        # Extract contact information
        contact_info = extract_contact_info(processed_text[:1000])  # Usually at the beginning
        parsed_data["contact_info"] = contact_info
        
        # Extract education
        if "education" in sections:
            education = extract_education(sections["education"])
            parsed_data["education"] = education
        
        # Extract experience
        if "experience" in sections:
            experience = extract_experience(sections["experience"])
            parsed_data["experience"] = experience
        
        # Extract skills
        skills = []
        if "skills" in sections:
            skills = extract_skills(sections["skills"])
        else:
            # If no explicit skills section, look throughout the document
            skills = extract_skills(processed_text)
        
        parsed_data["skills"] = skills
        
        # Use NLP for additional entity extraction
        entities = extract_entities_with_nlp(processed_text)
        
        # Merge NLP-extracted skills with pattern-matched skills
        all_skills = list(set(skills + entities.get("skills", [])))
        entities["skills"] = all_skills
        
        parsed_data["entities"] = entities
        
    except Exception as e:
        logger.error(f"Error parsing resume: {e}")
        parsed_data["error"] = str(e)
    
    return parsed_data


# Standalone functions for direct usage

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from a file regardless of format
    
    Args:
        file_path: Path to file
        
    Returns:
        str: Extracted text
    """
    return extract_text_from_resume(file_path)


def extract_skills_from_text(text: str) -> List[str]:
    """
    Extract skills from text
    
    Args:
        text: Text to extract skills from
        
    Returns:
        list: Extracted skills
    """
    return extract_skills(text)


def extract_experience_from_text(text: str) -> List[Dict[str, str]]:
    """
    Extract work experience from text
    
    Args:
        text: Text to extract experience from
        
    Returns:
        list: Extracted experience entries
    """
    return extract_experience(text)


def get_resume_summary(file_path: str) -> Dict[str, Any]:
    """
    Get a summary of a resume
    
    Args:
        file_path: Path to resume file
        
    Returns:
        dict: Resume summary
    """
    parsed_resume = parse_resume(file_path)
    
    # Create a simplified summary
    summary = {
        "contact_info": parsed_resume.get("contact_info", {}),
        "skills": parsed_resume.get("skills", []),
        "education_count": len(parsed_resume.get("education", [])),
        "experience_count": len(parsed_resume.get("experience", [])),
        "word_count": parsed_resume.get("metadata", {}).get("word_count", 0),
        "file_type": parsed_resume.get("file_type", ""),
        "has_error": "error" in parsed_resume
    }
    
    # Add most recent experience
    experiences = parsed_resume.get("experience", [])
    if experiences:
        # Sort by start date if available
        sorted_exp = sorted(experiences, 
                          key=lambda x: x.get("date_range", "").split()[0], 
                          reverse=True)
        summary["latest_job"] = sorted_exp[0].get("job_title", "")
        summary["latest_company"] = sorted_exp[0].get("company", "")
    
    # Add highest education
    education = parsed_resume.get("education", [])
    if education:
        degrees = [edu.get("degree", "") for edu in education]
        degree_levels = {
            "phd": 5,
            "doctor": 5,
            "doctorate": 5,
            "master": 4,
            "mba": 4,
            "ms": 4,
            "ma": 4,
            "bachelor": 3,
            "bs": 3,
            "ba": 3,
            "associate": 2,
            "diploma": 1,
            "certificate": 0
        }
        
        highest_level = -1
        highest_degree = ""
        
        for degree in degrees:
            for level_name, level_value in degree_levels.items():
                if level_name in degree.lower() and level_value > highest_level:
                    highest_level = level_value
                    highest_degree = degree
        
        if highest_degree:
            summary["highest_education"] = highest_degree
    
    return summary