import os
import json
import logging
import uuid
import tempfile
from typing import Dict, List, Tuple, Any, Optional, Union
from datetime import datetime
import re
import base64

# Optional dependencies - allow graceful fallback if not available
try:
    import jinja2
    JINJA2_AVAILABLE = True
except ImportError:
    JINJA2_AVAILABLE = False

try:
    import pdfkit
    PDFKIT_AVAILABLE = True
except ImportError:
    PDFKIT_AVAILABLE = False

try:
    import weasyprint
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ResumeBuilder:
    """
    Helps users build professional resumes using templates and AI-powered
    content suggestions. Supports multiple formats and ATS-friendly layouts.
    """
    
    def __init__(self, 
               template_dir: Optional[str] = None,
               output_dir: Optional[str] = None,
               default_template: str = "modern",
               pdf_engine: str = "auto"):
        """
        Initialize the resume builder
        
        Args:
            template_dir: Directory containing resume templates
            output_dir: Directory for resume outputs
            default_template: Default template to use
            pdf_engine: PDF generation engine ("weasyprint", "pdfkit", or "auto")
        """
        self.logger = logging.getLogger(__name__)
        self.default_template = default_template
        self.pdf_engine = pdf_engine
        
        # Set up template directory
        self.template_dir = template_dir
        if not self.template_dir or not os.path.exists(self.template_dir):
            # Look for templates in standard locations
            possible_dirs = [
                os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates", "resume"),
                os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "templates", "resume"),
                os.path.join(os.path.expanduser("~"), ".tamkeen", "templates", "resume")
            ]
            
            for dir_path in possible_dirs:
                if os.path.exists(dir_path):
                    self.template_dir = dir_path
                    break
        
        # Set up output directory
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            self.output_dir = output_dir
        else:
            self.output_dir = tempfile.gettempdir()
            
        # Initialize template engine
        self.jinja_env = None
        if JINJA2_AVAILABLE and self.template_dir:
            try:
                self.jinja_env = jinja2.Environment(
                    loader=jinja2.FileSystemLoader(self.template_dir),
                    autoescape=jinja2.select_autoescape(['html', 'xml'])
                )
                
                # Add custom filters
                self.jinja_env.filters['date_format'] = self._date_format_filter
                self.jinja_env.filters['safe_html'] = self._safe_html_filter
            except Exception as e:
                self.logger.error(f"Error initializing Jinja2: {str(e)}")
        
        # Set up available templates
        self.available_templates = self._get_available_templates()
        
        # Set up available output formats
        self.available_formats = ["html"]
        if PDFKIT_AVAILABLE or WEASYPRINT_AVAILABLE:
            self.available_formats.append("pdf")
        if DOCX_AVAILABLE:
            self.available_formats.append("docx")
        self.available_formats.append("json")  # Always available
        
        # Define standard sections
        self.standard_sections = [
            "contact", "summary", "experience", "education", 
            "skills", "certifications", "projects", "awards",
            "publications", "languages", "volunteer", "interests"
        ]
        
        self.logger.info(f"Resume builder initialized with {len(self.available_templates)} templates")
    
    def _get_available_templates(self) -> Dict[str, Dict[str, Any]]:
        """Get list of available resume templates"""
        templates = {}
        
        # Built-in templates
        built_in = {
            "modern": {
                "name": "Modern",
                "description": "Clean and professional template with a modern layout",
                "preview": None,
                "formats": ["html", "pdf"],
                "ats_friendly": True
            },
            "classic": {
                "name": "Classic",
                "description": "Traditional resume format with a timeless design",
                "preview": None,
                "formats": ["html", "pdf", "docx"],
                "ats_friendly": True
            },
            "creative": {
                "name": "Creative",
                "description": "Bold design for creative professionals",
                "preview": None,
                "formats": ["html", "pdf"],
                "ats_friendly": False
            },
            "minimalist": {
                "name": "Minimalist",
                "description": "Clean, simple design with focus on content",
                "preview": None,
                "formats": ["html", "pdf", "docx"],
                "ats_friendly": True
            },
            "executive": {
                "name": "Executive",
                "description": "Sophisticated design for senior professionals",
                "preview": None,
                "formats": ["html", "pdf", "docx"],
                "ats_friendly": True
            }
        }
        
        # Add built-in templates
        templates.update(built_in)
        
        # Check template directory for additional templates
        if self.template_dir and os.path.exists(self.template_dir):
            for item in os.listdir(self.template_dir):
                if not item.startswith('_') and not item.startswith('.'):
                    # Check if it's a directory or an HTML file
                    if os.path.isdir(os.path.join(self.template_dir, item)):
                        template_id = item
                    elif item.endswith('.html') or item.endswith('.jinja2'):
                        template_id = os.path.splitext(item)[0]
                    else:
                        continue
                        
                    # Only add if not already in built-in templates
                    if template_id not in templates:
                        # Try to load metadata
                        metadata = self._load_template_metadata(template_id)
                        if metadata:
                            templates[template_id] = metadata
                        else:
                            # Default metadata
                            templates[template_id] = {
                                "name": template_id.replace('_', ' ').title(),
                                "description": f"Custom template: {template_id}",
                                "preview": None,
                                "formats": ["html", "pdf"] if PDFKIT_AVAILABLE or WEASYPRINT_AVAILABLE else ["html"],
                                "ats_friendly": True
                            }
        
        return templates
    
    def _load_template_metadata(self, template_id: str) -> Optional[Dict[str, Any]]:
        """Load metadata for a specific template"""
        metadata_file = os.path.join(self.template_dir, template_id, "metadata.json")
        if os.path.exists(metadata_file):
            try:
                with open(metadata_file, 'r') as f:
                    return json.load(f)
            except Exception as e:
                self.logger.error(f"Error loading template metadata for {template_id}: {str(e)}")
        return None
    
    def _date_format_filter(self, value, format_str="%B %Y"):
        """Jinja2 filter to format dates"""
        if not value:
            return ""
            
        try:
            if isinstance(value, str):
                # Try to parse the date string
                formats = [
                    "%Y-%m-%d", "%Y/%m/%d", "%d-%m-%Y", "%d/%m/%Y",
                    "%Y-%m", "%Y/%m", "%m-%Y", "%m/%Y",
                    "%B %Y", "%b %Y"
                ]
                
                for fmt in formats:
                    try:
                        date_obj = datetime.strptime(value, fmt)
                        return date_obj.strftime(format_str)
                    except ValueError:
                        continue
                        
                # If no format matched, return as is
                return value
            elif isinstance(value, datetime):
                return value.strftime(format_str)
            else:
                return str(value)
        except Exception:
            return str(value)
    
    def _safe_html_filter(self, value):
        """Make HTML safe for output in templates"""
        if not value:
            return ""
            
        # Basic HTML escaping
        value = str(value)
        value = value.replace('&', '&amp;')
        value = value.replace('<', '&lt;')
        value = value.replace('>', '&gt;')
        value = value.replace('"', '&quot;')
        value = value.replace("'", '&#39;')
        
        # Convert newlines to <br>
        value = value.replace('\n', '<br>')
        
        return value
    
    def get_templates(self) -> Dict[str, Dict[str, Any]]:
        """Get list of available templates with metadata"""
        return self.available_templates
    
    def get_template_preview(self, template_id: str) -> Optional[str]:
        """Get preview image for a template as base64 string"""
        if template_id not in self.available_templates:
            return None
            
        # Check for preview image
        preview_paths = [
            os.path.join(self.template_dir, template_id, "preview.png"),
            os.path.join(self.template_dir, template_id, "preview.jpg"),
            os.path.join(self.template_dir, f"{template_id}_preview.png"),
            os.path.join(self.template_dir, f"{template_id}_preview.jpg")
        ]
        
        for path in preview_paths:
            if os.path.exists(path):
                try:
                    with open(path, 'rb') as f:
                        image_data = f.read()
                        image_type = path.split('.')[-1]
                        return f"data:image/{image_type};base64,{base64.b64encode(image_data).decode('utf-8')}"
                except Exception as e:
                    self.logger.error(f"Error loading template preview: {str(e)}")
                    
        return None
    
    def create_resume(self, 
                    resume_data: Dict[str, Any],
                    template_id: Optional[str] = None,
                    output_format: str = "pdf",
                    filename: Optional[str] = None) -> Dict[str, Any]:
        """
        Create a resume from provided data
        
        Args:
            resume_data: Resume content data
            template_id: Template to use
            output_format: Output format (pdf, html, docx, json)
            filename: Optional output filename
            
        Returns:
            Dictionary with result information
        """
        # Validate and prepare resume data
        processed_data = self._prepare_resume_data(resume_data)
        
        # Select template
        template_id = template_id or self.default_template
        if template_id not in self.available_templates:
            self.logger.warning(f"Template {template_id} not found, using {self.default_template}")
            template_id = self.default_template
            
        # Validate output format
        output_format = output_format.lower()
        if output_format not in self.available_formats:
            self.logger.warning(f"Format {output_format} not available, using html")
            output_format = "html"
            
        # Generate output filename if not provided
        if not filename:
            name_part = re.sub(r'[^a-zA-Z0-9]', '_', processed_data.get("contact", {}).get("name", "resume")).lower()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{name_part}_{timestamp}.{output_format}"
            filename = os.path.join(self.output_dir, filename)
        
        # Create directory if it doesn't exist
        os.makedirs(os.path.dirname(os.path.abspath(filename)), exist_ok=True)
        
        # Generate resume based on format
        result = {
            "success": False,
            "filename": filename,
            "format": output_format,
            "template": template_id,
            "message": ""
        }
        
        try:
            if output_format == "json":
                # JSON output is just the processed data
                with open(filename, 'w') as f:
                    json.dump(processed_data, f, indent=2)
                result["success"] = True
                result["message"] = "Resume data saved as JSON"
                
            elif output_format == "html":
                # Generate HTML
                html_content = self._generate_html_resume(processed_data, template_id)
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                result["success"] = True
                result["message"] = "HTML resume created successfully"
                result["content"] = html_content
                
            elif output_format == "pdf":
                # Generate PDF from HTML
                html_content = self._generate_html_resume(processed_data, template_id)
                pdf_created = self._generate_pdf(html_content, filename)
                
                if pdf_created:
                    result["success"] = True
                    result["message"] = "PDF resume created successfully"
                else:
                    result["success"] = False
                    result["message"] = "Failed to generate PDF resume"
                    
            elif output_format == "docx":
                # Generate DOCX
                if DOCX_AVAILABLE:
                    docx_created = self._generate_docx_resume(processed_data, template_id, filename)
                    if docx_created:
                        result["success"] = True
                        result["message"] = "DOCX resume created successfully"
                    else:
                        result["success"] = False
                        result["message"] = "Failed to generate DOCX resume"
                else:
                    result["success"] = False
                    result["message"] = "DOCX generation not available (python-docx not installed)"
        
        except Exception as e:
            self.logger.error(f"Error generating resume: {str(e)}")
            result["success"] = False
            result["message"] = f"Error: {str(e)}"
            
        return result
    
    def _prepare_resume_data(self, resume_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate and prepare resume data"""
        # Create a copy to avoid modifying the original
        processed_data = resume_data.copy()
        
        # Ensure all standard sections exist
        for section in self.standard_sections:
            if section not in processed_data:
                if section in ["experience", "education", "skills", "certifications", "projects"]:
                    processed_data[section] = []
                else:
                    processed_data[section] = {}
        
        # Ensure contact section has required fields
        if "contact" not in processed_data:
            processed_data["contact"] = {}
            
        contact = processed_data["contact"]
        required_contact_fields = ["name", "email", "phone"]
        for field in required_contact_fields:
            if field not in contact:
                contact[field] = ""
                
        # Generate a timestamp for the resume
        processed_data["generated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Add a resume ID if not present
        if "resume_id" not in processed_data:
            processed_data["resume_id"] = str(uuid.uuid4())
            
        return processed_data
    
    def _generate_html_resume(self, resume_data: Dict[str, Any], template_id: str) -> str:
        """Generate HTML resume using the selected template"""
        if not JINJA2_AVAILABLE or not self.jinja_env:
            # Fallback to basic HTML if Jinja2 not available
            return self._generate_basic_html_resume(resume_data)
            
        try:
            # Look for template file
            template_paths = [
                f"{template_id}.html",
                f"{template_id}.jinja2",
                f"{template_id}/index.html",
                f"{template_id}/template.html",
                f"{template_id}/resume.html",
                f"{template_id}/resume.jinja2"
            ]
            
            template = None
            for path in template_paths:
                try:
                    template = self.jinja_env.get_template(path)
                    break
                except jinja2.exceptions.TemplateNotFound:
                    continue
                    
            if not template:
                self.logger.warning(f"Template {template_id} not found, using basic HTML")
                return self._generate_basic_html_resume(resume_data)
                
            # Render template with resume data
            html_content = template.render(**resume_data)
            return html_content
            
        except Exception as e:
            self.logger.error(f"Error generating HTML resume: {str(e)}")
            return self._generate_basic_html_resume(resume_data)
    
    def _generate_basic_html_resume(self, resume_data: Dict[str, Any]) -> str:
        """Generate a basic HTML resume without templates"""
        # Simple HTML template for fallback
        html = [
            "<!DOCTYPE html>",
            "<html lang='en'>",
            "<head>",
            "    <meta charset='UTF-8'>",
            "    <meta name='viewport' content='width=device-width, initial-scale=1.0'>",
            f"    <title>{resume_data.get('contact', {}).get('name', 'Resume')}</title>",
            "    <style>",
            "        body { font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }",
            "        h1 { color: #2c3e50; }",
            "        h2 { color: #3498db; border-bottom: 1px solid #eee; padding-bottom: 5px; }",
            "        .contact { margin-bottom: 20px; }",
            "        .section { margin-bottom: 30px; }",
            "        .item { margin-bottom: 20px; }",
            "        .item-header { display: flex; justify-content: space-between; }",
            "        .item-title { font-weight: bold; }",
            "        .item-location, .item-date { color: #7f8c8d; }",
            "        .skills { display: flex; flex-wrap: wrap; }",
            "        .skill { background: #eee; padding: 5px 10px; margin-right: 5px; margin-bottom: 5px; border-radius: 3px; }",
            "    </style>",
            "</head>",
            "<body>"
        ]
        
        # Contact section
        contact = resume_data.get("contact", {})
        html.append(f"<h1>{contact.get('name', '')}</h1>")
        html.append("<div class='contact'>")
        if contact.get("email"):
            html.append(f"<div>{contact.get('email')}</div>")
        if contact.get("phone"):
            html.append(f"<div>{contact.get('phone')}</div>")
        if contact.get("address"):
            html.append(f"<div>{contact.get('address')}</div>")
        if contact.get("linkedin"):
            html.append(f"<div>LinkedIn: {contact.get('linkedin')}</div>")
        if contact.get("website"):
            html.append(f"<div>Website: {contact.get('website')}</div>")
        html.append("</div>")
        
        # Summary section
        if resume_data.get("summary"):
            html.append("<div class='section'>")
            html.append("<h2>Summary</h2>")
            html.append(f"<p>{resume_data.get('summary')}</p>")
            html.append("</div>")
        
        # Experience section
        experiences = resume_data.get("experience", [])
        if experiences:
            html.append("<div class='section'>")
            html.append("<h2>Experience</h2>")
            for exp in experiences:
                html.append("<div class='item'>")
                html.append("<div class='item-header'>")
                html.append(f"<span class='item-title'>{exp.get('title', '')}</span>")
                html.append(f"<span class='item-date'>{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}</span>")
                html.append("</div>")
                html.append(f"<div class='item-location'>{exp.get('company', '')}, {exp.get('location', '')}</div>")
                html.append(f"<div class='item-description'>{exp.get('description', '').replace(chr(10), '<br>')}</div>")
                html.append("</div>")
            html.append("</div>")
        
        # Education section
        education = resume_data.get("education", [])
        if education:
            html.append("<div class='section'>")
            html.append("<h2>Education</h2>")
            for edu in education:
                html.append("<div class='item'>")
                html.append("<div class='item-header'>")
                html.append(f"<span class='item-title'>{edu.get('degree', '')}, {edu.get('field_of_study', '')}</span>")
                html.append(f"<span class='item-date'>{edu.get('start_date', '')} - {edu.get('end_date', 'Present')}</span>")
                html.append("</div>")
                html.append(f"<div class='item-location'>{edu.get('institution', '')}, {edu.get('location', '')}</div>")
                if edu.get("description"):
                    html.append(f"<div class='item-description'>{edu.get('description', '')}</div>")
                html.append("</div>")
            html.append("</div>")
        
        # Skills section
        skills = resume_data.get("skills", [])
        if skills:
            html.append("<div class='section'>")
            html.append("<h2>Skills</h2>")
            html.append("<div class='skills'>")
            for skill in skills:
                if isinstance(skill, str):
                    html.append(f"<div class='skill'>{skill}</div>")
                elif isinstance(skill, dict) and "name" in skill:
                    html.append(f"<div class='skill'>{skill['name']}</div>")
            html.append("</div>")
            html.append("</div>")
        
        # Projects section
        projects = resume_data.get("projects", [])
        if projects:
            html.append("<div class='section'>")
            html.append("<h2>Projects</h2>")
            for project in projects:
                html.append("<div class='item'>")
                html.append("<div class='item-header'>")
                html.append(f"<span class='item-title'>{project.get('name', '')}</span>")
                if project.get("date"):
                    html.append(f"<span class='item-date'>{project.get('date', '')}</span>")
                html.append("</div>")
                if project.get("description"):
                    html.append(f"<div class='item-description'>{project.get('description', '')}</div>")
                html.append("</div>")
            html.append("</div>")
        
        # Certifications section
        certifications = resume_data.get("certifications", [])
        if certifications:
            html.append("<div class='section'>")
            html.append("<h2>Certifications</h2>")
            for cert in certifications:
                html.append("<div class='item'>")
                html.append("<div class='item-header'>")
                html.append(f"<span class='item-title'>{cert.get('name', '')}</span>")
                if cert.get("date"):
                    html.append(f"<span class='item-date'>{cert.get('date', '')}</span>")
                html.append("</div>")
                if cert.get("issuer"):
                    html.append(f"<div class='item-location'>{cert.get('issuer', '')}</div>")
                html.append("</div>")
            html.append("</div>")
            
        # Close HTML
        html.append("</body>")
        html.append("</html>")
        
        return "\n".join(html)
    
    def _generate_pdf(self, html_content: str, output_file: str) -> bool:
        """Generate PDF from HTML content"""
        if not html_content:
            return False
            
        # Determine which PDF engine to use
        if self.pdf_engine == "weasyprint" or (self.pdf_engine == "auto" and WEASYPRINT_AVAILABLE):
            try:
                if WEASYPRINT_AVAILABLE:
                    import weasyprint
                    weasyprint.HTML(string=html_content).write_pdf(output_file)
                    return True
                else:
                    self.logger.error("WeasyPrint not available")
                    return False
            except Exception as e:
                self.logger.error(f"Error generating PDF with WeasyPrint: {str(e)}")
                return False
                
        elif self.pdf_engine == "pdfkit" or (self.pdf_engine == "auto" and PDFKIT_AVAILABLE):
            try:
                if PDFKIT_AVAILABLE:
                    import pdfkit
                    # Use wkhtmltopdf options for better rendering
                    options = {
                        'page-size': 'Letter',
                        'margin-top': '0.75in',
                        'margin-right': '0.75in',
                        'margin-bottom': '0.75in',
                        'margin-left': '0.75in',
                        'encoding': 'UTF-8',
                        'no-outline': None,
                        'enable-local-file-access': None
                    }
                    pdfkit.from_string(html_content, output_file, options=options)
                    return True
                else:
                    self.logger.error("PDFKit not available")
                    return False
            except Exception as e:
                self.logger.error(f"Error generating PDF with PDFKit: {str(e)}")
                return False
        
        self.logger.error("No PDF generation engine available")
        return False
    
    def _generate_docx_resume(self, resume_data: Dict[str, Any], template_id: str, output_file: str) -> bool:
        """Generate DOCX format resume"""
        if not DOCX_AVAILABLE:
            return False
            
        try:
            # Create new Document
            doc = docx.Document()
            
            # Set document properties
            doc.core_properties.author = resume_data.get("contact", {}).get("name", "")
            doc.core_properties.title = f"Resume - {resume_data.get('contact', {}).get('name', '')}"
            
            # Apply template style
            self._apply_docx_template_style(doc, template_id)
            
            # Add contact information
            contact = resume_data.get("contact", {})
            name = contact.get("name", "")
            if name:
                heading = doc.add_paragraph()
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                heading_run = heading.add_run(name)
                heading_run.bold = True
                heading_run.font.size = Pt(18)
                
            # Contact details
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            contact_items = []
            if contact.get("email"):
                contact_items.append(contact.get("email"))
            if contact.get("phone"):
                contact_items.append(contact.get("phone"))
            if contact.get("address"):
                contact_items.append(contact.get("address"))
            if contact.get("linkedin"):
                contact_items.append(contact.get("linkedin"))
            
            contact_text = " | ".join(contact_items)
            contact_para.add_run(contact_text)
            
            # Add summary
            if resume_data.get("summary"):
                doc.add_paragraph()  # Spacer
                doc.add_heading("Professional Summary", level=1)
                summary_para = doc.add_paragraph()
                summary_para.add_run(resume_data.get("summary"))
            
            # Add experience
            experiences = resume_data.get("experience", [])
            if experiences:
                doc.add_paragraph()  # Spacer
                doc.add_heading("Professional Experience", level=1)
                
                for exp in experiences:
                    # Job title and company
                    title_para = doc.add_paragraph()
                    title_run = title_para.add_run(exp.get("title", ""))
                    title_run.bold = True
                    
                    company_para = doc.add_paragraph()
                    company_text = f"{exp.get('company', '')}"
                    if exp.get("location"):
                        company_text += f", {exp.get('location', '')}"
                    company_run = company_para.add_run(company_text)
                    company_run.italic = True
                    
                    # Dates
                    date_para = doc.add_paragraph()
                    date_text = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}"
                    date_para.add_run(date_text)
                    
                    # Description
                    if exp.get("description"):
                        desc = exp.get("description", "")
                        # Handle bullet points if description has newlines
                        if "\n" in desc:
                            lines = desc.split("\n")
                            for line in lines:
                                if line.strip():
                                    bullet_para = doc.add_paragraph(style="List Bullet")
                                    bullet_para.add_run(line.strip())
                        else:
                            doc.add_paragraph(desc)
                    
                    doc.add_paragraph()  # Spacer
            
            # Add education
            education = resume_data.get("education", [])
            if education:
                doc.add_paragraph()  # Spacer
                doc.add_heading("Education", level=1)
                
                for edu in education:
                    # Degree and institution
                    degree_para = doc.add_paragraph()
                    degree_text = f"{edu.get('degree', '')}"
                    if edu.get("field_of_study"):
                        degree_text += f", {edu.get('field_of_study', '')}"
                    degree_run = degree_para.add_run(degree_text)
                    degree_run.bold = True
                    
                    institution_para = doc.add_paragraph()
                    institution_text = f"{edu.get('institution', '')}"
                    if edu.get("location"):
                        institution_text += f", {edu.get('location', '')}"
                    institution_run = institution_para.add_run(institution_text)
                    institution_run.italic = True
                    
                    # Dates
                    date_para = doc.add_paragraph()
                    date_text = f"{edu.get('start_date', '')} - {edu.get('end_date', 'Present')}"
                    date_para.add_run(date_text)
                    
                    # Description
                    if edu.get("description"):
                        doc.add_paragraph(edu.get("description"))
                    
                    doc.add_paragraph()  # Spacer
            
            # Add skills
            skills = resume_data.get("skills", [])
            if skills:
                doc.add_paragraph()  # Spacer
                doc.add_heading("Skills", level=1)
                
                skills_para = doc.add_paragraph()
                skill_names = []
                
                for skill in skills:
                    if isinstance(skill, str):
                        skill_names.append(skill)
                    elif isinstance(skill, dict) and "name" in skill:
                        skill_names.append(skill["name"])
                
                skills_para.add_run(", ".join(skill_names))
            
            # Add projects
            projects = resume_data.get("projects", [])
            if projects:
                doc.add_paragraph()  # Spacer
                doc.add_heading("Projects", level=1)
                
                for project in projects:
                    # Project name
                    project_para = doc.add_paragraph()
                    project_run = project_para.add_run(project.get("name", ""))
                    project_run.bold = True
                    
                    # Date
                    if project.get("date"):
                        date_para = doc.add_paragraph()
                        date_para.add_run(project.get("date"))
                    
                    # Description
                    if project.get("description"):
                        doc.add_paragraph(project.get("description"))
                    
                    doc.add_paragraph()  # Spacer
            
            # Save document
            doc.save(output_file)
            return True
            
        except Exception as e:
            self.logger.error(f"Error generating DOCX resume: {str(e)}")
            return False
    
    def _apply_docx_template_style(self, doc, template_id: str):
        """Apply template-specific styling to DOCX document"""
        # Set default styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Apply template-specific styles
        if template_id == "classic":
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            
            # Heading styles
            heading1 = doc.styles['Heading 1']
            heading1.font.name = 'Times New Roman'
            heading1.font.size = Pt(14)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(0, 0, 0)
            
        elif template_id == "modern":
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            
            # Heading styles
            heading1 = doc.styles['Heading 1']
            heading1.font.name = 'Calibri'
            heading1.font.size = Pt(14)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(41, 128, 185)  # Blue
            
        elif template_id == "minimalist":
            style.font.name = 'Arial'
            style.font.size = Pt(10)
            
            # Heading styles
            heading1 = doc.styles['Heading 1']
            heading1.font.name = 'Arial'
            heading1.font.size = Pt(12)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(70, 70, 70)  # Dark gray
            
        elif template_id == "executive":
            style.font.name = 'Georgia'
            style.font.size = Pt(11)
            
            # Heading styles
            heading1 = doc.styles['Heading 1']
            heading1.font.name = 'Georgia'
            heading1.font.size = Pt(15)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
            
        # Set up list bullet style
        if 'List Bullet' not in doc.styles:
            list_style = doc.styles.add_style('List Bullet', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
            list_style.base_style = doc.styles['Normal']
            list_style.font.name = style.font.name
            list_style.font.size = style.font.size
        else:
            list_style = doc.styles['List Bullet']
            
        # Set paragraph spacing
        style.paragraph_format.space_after = Pt(6)
        list_style.paragraph_format.left_indent = Inches(0.25)
        
        # Add section break
        section = doc.sections[0]
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0) 