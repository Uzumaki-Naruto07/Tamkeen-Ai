"""
File Handler Module

This module provides utilities for handling, processing, and extracting information from files.
"""

import os
import re
import json
import uuid
from typing import Dict, List, Any, Optional, Tuple, Union
import logging
from datetime import datetime
import mimetypes

# Import settings
from backend.config.settings import UPLOAD_FOLDER, ALLOWED_EXTENSIONS, STORAGE_TYPE, STORAGE_CONFIG

# Setup logger
logger = logging.getLogger(__name__)

# Try to import optional dependencies
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Install with: pip install python-docx")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logger.warning("PyPDF2 not installed. Install with: pip install PyPDF2")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logger.warning("pdfplumber not installed. Install with: pip install pdfplumber")

try:
    from odf import text as odf_text
    from odf.opendocument import load as odf_load
    ODF_AVAILABLE = True
except ImportError:
    ODF_AVAILABLE = False
    logger.warning("odfpy not installed. Install with: pip install odfpy")

try:
    import boto3
    AWS_AVAILABLE = True
except ImportError:
    AWS_AVAILABLE = False
    logger.warning("boto3 not installed. Install with: pip install boto3")

try:
    from azure.storage.blob import BlobServiceClient
    AZURE_AVAILABLE = True
except ImportError:
    AZURE_AVAILABLE = False
    logger.warning("azure-storage-blob not installed. Install with: pip install azure-storage-blob")


def get_file_extension(filename: str) -> str:
    """
    Get file extension from filename
    
    Args:
        filename: Filename
        
    Returns:
        str: File extension without dot
    """
    return os.path.splitext(filename)[1][1:].lower()


def allowed_file(filename: str, file_type: str = 'document') -> bool:
    """
    Check if file has allowed extension
    
    Args:
        filename: Filename
        file_type: File type category (document, image, resume)
        
    Returns:
        bool: Is allowed
    """
    if not filename:
        return False
    
    extension = get_file_extension(filename)
    
    if file_type not in ALLOWED_EXTENSIONS:
        return False
    
    return extension in ALLOWED_EXTENSIONS[file_type]


def get_mime_type(filename: str) -> str:
    """
    Get MIME type for file
    
    Args:
        filename: Filename
        
    Returns:
        str: MIME type
    """
    mime_type, _ = mimetypes.guess_type(filename)
    return mime_type or 'application/octet-stream'


def generate_unique_filename(filename: str) -> str:
    """
    Generate unique filename
    
    Args:
        filename: Original filename
        
    Returns:
        str: Unique filename
    """
    extension = get_file_extension(filename)
    unique_id = str(uuid.uuid4())
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    new_filename = f"{unique_id}_{timestamp}.{extension}"
    
    return new_filename


def save_uploaded_file(file, folder_path: str = None, allowed_types: set = None) -> Optional[str]:
    """
    Save uploaded file to local storage
    
    Args:
        file: File object
        folder_path: Optional sub-folder path
        allowed_types: Optional set of allowed file extensions
        
    Returns:
        str or None: File path or None if save failed
    """
    try:
        if not file:
            return None
        
        # Check filename
        if not file.filename:
            return None
        
        # Check file extension if allowed_types provided
        if allowed_types:
            extension = get_file_extension(file.filename)
            if extension not in allowed_types:
                logger.warning(f"Invalid file type: {extension}")
                return None
        
        # Generate unique filename
        filename = generate_unique_filename(file.filename)
        
        # Determine full path
        if folder_path:
            full_folder_path = os.path.join(UPLOAD_FOLDER, folder_path)
        else:
            full_folder_path = UPLOAD_FOLDER
        
        # Ensure directory exists
        os.makedirs(full_folder_path, exist_ok=True)
        
        # Full file path
        file_path = os.path.join(full_folder_path, filename)
        
        # Save file
        file.save(file_path)
        
        logger.info(f"File saved: {file_path}")
        
        return file_path
    
    except Exception as e:
        logger.error(f"Error saving file: {str(e)}")
        return None


def upload_to_cloud_storage(file_path: str, destination_path: str = None) -> Optional[str]:
    """
    Upload file to cloud storage
    
    Args:
        file_path: Local file path
        destination_path: Optional destination path in storage
        
    Returns:
        str or None: Storage URL or None if upload failed
    """
    try:
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        filename = os.path.basename(file_path)
        
        if not destination_path:
            destination_path = filename
        
        if STORAGE_TYPE == 's3':
            return _upload_to_s3(file_path, destination_path)
        elif STORAGE_TYPE == 'azure':
            return _upload_to_azure(file_path, destination_path)
        else:
            # Local storage, just return the file path
            return file_path
    
    except Exception as e:
        logger.error(f"Error uploading to cloud storage: {str(e)}")
        return None


def _upload_to_s3(file_path: str, destination_path: str) -> Optional[str]:
    """Upload file to S3"""
    if not AWS_AVAILABLE:
        logger.error("boto3 not installed. Cannot upload to S3.")
        return None
    
    try:
        config = STORAGE_CONFIG.get('s3', {})
        bucket = config.get('bucket')
        
        if not bucket:
            logger.error("S3 bucket not configured")
            return None
        
        # Initialize S3 client
        s3_client = boto3.client(
            's3',
            region_name=config.get('region'),
            aws_access_key_id=config.get('access_key'),
            aws_secret_access_key=config.get('secret_key'),
            endpoint_url=config.get('endpoint_url')
        )
        
        # Upload file
        s3_client.upload_file(file_path, bucket, destination_path)
        
        # Generate URL
        url = f"https://{bucket}.s3.amazonaws.com/{destination_path}"
        
        logger.info(f"File uploaded to S3: {url}")
        
        return url
    
    except Exception as e:
        logger.error(f"Error uploading to S3: {str(e)}")
        return None


def _upload_to_azure(file_path: str, destination_path: str) -> Optional[str]:
    """Upload file to Azure Blob Storage"""
    if not AZURE_AVAILABLE:
        logger.error("azure-storage-blob not installed. Cannot upload to Azure.")
        return None
    
    try:
        config = STORAGE_CONFIG.get('azure', {})
        connection_string = config.get('connection_string')
        container_name = config.get('container')
        
        if not connection_string or not container_name:
            logger.error("Azure Storage not configured properly")
            return None
        
        # Initialize Azure client
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        container_client = blob_service_client.get_container_client(container_name)
        
        # Upload file
        with open(file_path, "rb") as data:
            blob_client = container_client.upload_blob(destination_path, data, overwrite=True)
        
        # Generate URL
        url = f"{blob_service_client.url}/{container_name}/{destination_path}"
        
        logger.info(f"File uploaded to Azure: {url}")
        
        return url
    
    except Exception as e:
        logger.error(f"Error uploading to Azure: {str(e)}")
        return None


def delete_file(file_path: str) -> bool:
    """
    Delete file from storage
    
    Args:
        file_path: File path or URL
        
    Returns:
        bool: Success status
    """
    try:
        # If it's a local file
        if os.path.exists(file_path):
            os.remove(file_path)
            logger.info(f"File deleted: {file_path}")
            return True
        
        # If it's a cloud storage URL
        if STORAGE_TYPE == 's3' and file_path.startswith('https://'):
            return _delete_from_s3(file_path)
        elif STORAGE_TYPE == 'azure' and file_path.startswith('https://'):
            return _delete_from_azure(file_path)
        
        logger.warning(f"File not found: {file_path}")
        return False
    
    except Exception as e:
        logger.error(f"Error deleting file: {str(e)}")
        return False


def _delete_from_s3(file_url: str) -> bool:
    """Delete file from S3"""
    if not AWS_AVAILABLE:
        logger.error("boto3 not installed. Cannot delete from S3.")
        return False
    
    try:
        config = STORAGE_CONFIG.get('s3', {})
        bucket = config.get('bucket')
        
        if not bucket:
            logger.error("S3 bucket not configured")
            return False
        
        # Extract object key from URL
        # URL format: https://bucket-name.s3.amazonaws.com/object-key
        object_key = file_url.split(f"{bucket}.s3.amazonaws.com/")[1]
        
        # Initialize S3 client
        s3_client = boto3.client(
            's3',
            region_name=config.get('region'),
            aws_access_key_id=config.get('access_key'),
            aws_secret_access_key=config.get('secret_key'),
            endpoint_url=config.get('endpoint_url')
        )
        
        # Delete file
        s3_client.delete_object(Bucket=bucket, Key=object_key)
        
        logger.info(f"File deleted from S3: {file_url}")
        
        return True
    
    except Exception as e:
        logger.error(f"Error deleting from S3: {str(e)}")
        return False


def _delete_from_azure(file_url: str) -> bool:
    """Delete file from Azure Blob Storage"""
    if not AZURE_AVAILABLE:
        logger.error("azure-storage-blob not installed. Cannot delete from Azure.")
        return False
    
    try:
        config = STORAGE_CONFIG.get('azure', {})
        connection_string = config.get('connection_string')
        container_name = config.get('container')
        
        if not connection_string or not container_name:
            logger.error("Azure Storage not configured properly")
            return False
        
        # Extract blob name from URL
        # URL format: https://account.blob.core.windows.net/container/blob-name
        blob_name = file_url.split(f"{container_name}/")[1]
        
        # Initialize Azure client
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        container_client = blob_service_client.get_container_client(container_name)
        
        # Delete blob
        container_client.delete_blob(blob_name)
        
        logger.info(f"File deleted from Azure: {file_url}")
        
        return True
    
    except Exception as e:
        logger.error(f"Error deleting from Azure: {str(e)}")
        return False


def get_file_url(file_path: str) -> str:
    """
    Get public URL for file
    
    Args:
        file_path: File path
        
    Returns:
        str: File URL
    """
    # If already a URL, return as is
    if file_path.startswith(('http://', 'https://')):
        return file_path
    
    # For local storage, determine relative path from UPLOAD_FOLDER
    if file_path.startswith(UPLOAD_FOLDER):
        rel_path = os.path.relpath(file_path, UPLOAD_FOLDER)
        return f"/uploads/{rel_path}"
    
    # Otherwise return as is
    return file_path


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from document file
    
    Args:
        file_path: File path
        
    Returns:
        str: Extracted text
    """
    try:
        extension = get_file_extension(file_path)
        
        if extension == 'txt':
            return _extract_text_from_txt(file_path)
        elif extension == 'pdf':
            return _extract_text_from_pdf(file_path)
        elif extension in ['doc', 'docx']:
            return _extract_text_from_docx(file_path)
        elif extension in ['odt']:
            return _extract_text_from_odt(file_path)
        elif extension == 'rtf':
            return _extract_text_from_rtf(file_path)
        else:
            logger.warning(f"Unsupported file format for text extraction: {extension}")
            return ""
    
    except Exception as e:
        logger.error(f"Error extracting text from file: {str(e)}")
        return ""


def _extract_text_from_txt(file_path: str) -> str:
    """Extract text from .txt file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    except Exception as e:
        logger.error(f"Error reading .txt file: {str(e)}")
        
        # Try another encoding
        try:
            with open(file_path, 'r', encoding='latin-1', errors='ignore') as file:
                return file.read()
        except Exception as e2:
            logger.error(f"Error reading .txt file with latin-1 encoding: {str(e2)}")
            return ""


def _extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    text = ""
    
    # Try using pdfplumber first (better quality)
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
                    text += "\n\n"
            
            if text.strip():
                return text
            # If text is empty, fall back to PyPDF2
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed, falling back to PyPDF2: {str(e)}")
    
    # Fall back to PyPDF2
    if PYPDF2_AVAILABLE:
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num in range(len(reader.pages)):
                    text += reader.pages[page_num].extract_text() or ""
                    text += "\n\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF with PyPDF2: {str(e)}")
    
    logger.error("No PDF extraction library available. Install PyPDF2 or pdfplumber.")
    return ""


def _extract_text_from_docx(file_path: str) -> str:
    """Extract text from .docx file"""
    if not DOCX_AVAILABLE:
        logger.error("python-docx not installed. Cannot extract text from .docx file.")
        return ""
    
    try:
        doc = docx.Document(file_path)
        text = []
        
        for para in doc.paragraphs:
            text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        
        return "\n".join(text)
    except Exception as e:
        logger.error(f"Error extracting text from .docx: {str(e)}")
        return ""


def _extract_text_from_odt(file_path: str) -> str:
    """Extract text from .odt file"""
    if not ODF_AVAILABLE:
        logger.error("odfpy not installed. Cannot extract text from .odt file.")
        return ""
    
    try:
        doc = odf_load(file_path)
        allparas = doc.getElementsByType(odf_text.P)
        text = []
        
        for para in allparas:
            text.append(odf_text.teletype.extractText(para))
        
        return "\n".join(text)
    except Exception as e:
        logger.error(f"Error extracting text from .odt: {str(e)}")
        return ""


def _extract_text_from_rtf(file_path: str) -> str:
    """
    Extract text from .rtf file
    
    Note: This is a very simple RTF parser. For complex RTF files,
    consider using a specialized library.
    """
    try:
        with open(file_path, 'r', encoding='latin-1', errors='ignore') as file:
            content = file.read()
        
        # Very basic RTF parsing - removing RTF markup
        # Remove RTF header
        text = re.sub(r'^.*\\rtf1\\.*?\\pard', '', content, flags=re.DOTALL)
        
        # Remove RTF commands
        text = re.sub(r'\\[a-z0-9]+', ' ', text)
        text = re.sub(r'\\\'[0-9a-f]{2}', '', text)
        text = re.sub(r'\\[^a-z]', '', text)
        text = re.sub(r'\{.*?\}', '', text, flags=re.DOTALL)
        
        # Replace RTF line breaks with normal line breaks
        text = re.sub(r'\\par\s*', '\n', text)
        
        # Remove any remaining braces
        text = text.replace('{', '').replace('}', '')
        
        return text
    except Exception as e:
        logger.error(f"Error extracting text from .rtf: {str(e)}")
        return ""


def detect_sections(text: str) -> Dict[str, str]:
    """
    Detect document sections
    
    Args:
        text: Document text
        
    Returns:
        dict: Detected sections
    """
    # Common resume section headings
    section_patterns = {
        'contact': r'(?i)(contact|personal)(?:\s+information)?',
        'summary': r'(?i)(summary|profile|objective|about me|personal statement)',
        'education': r'(?i)(education|academic|qualifications)',
        'experience': r'(?i)(experience|employment|work|history|background)',
        'skills': r'(?i)(skills|expertise|competencies|technologies)',
        'projects': r'(?i)(projects|portfolio)',
        'certifications': r'(?i)(certifications|certificates|licenses)',
        'languages': r'(?i)(languages)',
        'awards': r'(?i)(awards|honors|achievements)',
        'publications': r'(?i)(publications|research)',
        'references': r'(?i)(references|recommendations)',
        'volunteer': r'(?i)(volunteer|community)'
    }
    
    # Find all potential section headings and their positions
    section_positions = []
    for section, pattern in section_patterns.items():
        for match in re.finditer(pattern, text, re.MULTILINE):
            section_positions.append((match.start(), section, match.group()))
    
    # Sort by position
    section_positions.sort()
    
    if not section_positions:
        return {'full_text': text}
    
    # Extract text for each section
    sections = {}
    for i, (pos, section, heading) in enumerate(section_positions):
        start = pos
        
        # Find end position (start of next section or end of text)
        if i < len(section_positions) - 1:
            end = section_positions[i + 1][0]
        else:
            end = len(text)
        
        # Extract section text
        section_text = text[start:end].strip()
        sections[section] = section_text
    
    return sections
